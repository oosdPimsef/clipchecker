# -*- coding: utf-8 -*-

import tempfile
import unittest
from pathlib import Path

import docx

from app.document_text import extract_docx_text


class DocumentTextTests(unittest.TestCase):
    def test_extract_docx_text_includes_table_cells(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "booking_form.docx"
            document = docx.Document()
            document.add_paragraph("")
            table = document.add_table(rows=3, cols=2)
            table.cell(0, 0).text = "СВЕДЕНИЯ ОБ ИСПОЛЬЗОВАНИИ ПРОИЗВЕДЕНИЙ"
            table.cell(1, 0).text = "Название рекламного ролика"
            table.cell(1, 1).text = "Праздник SOKOLOV"
            table.cell(2, 0).text = "Продолжительность ролика"
            table.cell(2, 1).text = "5 сек."
            document.save(path)

            text = extract_docx_text(str(path))

        self.assertIn("СВЕДЕНИЯ ОБ ИСПОЛЬЗОВАНИИ ПРОИЗВЕДЕНИЙ", text)
        self.assertIn("Название рекламного ролика | Праздник SOKOLOV", text)
        self.assertIn("Продолжительность ролика | 5 сек.", text)


if __name__ == "__main__":
    unittest.main()
