# -*- coding: utf-8 -*-
"""Helpers for extracting full text from advertiser documents."""

from __future__ import annotations

import re

import docx


def clean_to_single_block(text: str) -> str:
    text = text.replace("\r", "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"(?<!\n)\n(?!\n)", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()


def extract_docx_text(filepath: str) -> str:
    document = docx.Document(filepath)
    parts: list[str] = []

    for para in document.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells: list[str] = []
            seen = set()
            for cell in row.cells:
                cell_text = clean_to_single_block(cell.text or "")
                if cell_text and cell_text not in seen:
                    cells.append(cell_text)
                    seen.add(cell_text)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)

